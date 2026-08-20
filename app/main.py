import io
import os
import re
import sys
import asyncio
from contextlib import asynccontextmanager
from html import escape
from pathlib import Path
from html.parser import HTMLParser
import base64

import httpx
from fastapi import FastAPI, HTTPException, Request, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from starlette.middleware.sessions import SessionMiddleware
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.auth import login_user, logout_user, get_user_from_session, templates, get_authenticated_client, build_display_name
from app.config import get_settings
from app.report_store import ensure_report_folders, save_dataframe, get_report_input_fields, get_report_folders
from app.report_data import build_tac2_dataframes, build_tac2_text_report

if sys.platform.startswith('win') and hasattr(asyncio, 'WindowsProactorEventLoopPolicy'):
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

BASE_DIR = Path(__file__).resolve().parents[1]
SETTINGS = get_settings()

FRONTEND_DIST = BASE_DIR / "frontend" / "dist"
PDF_BORDER_COLOR = "#cbd5e1"


def _patient_age_from_birth_date(birth_value):
    from datetime import date, datetime
    if not birth_value:
        return None
    try:
        birth_date = datetime.fromisoformat(str(birth_value)).date()
    except Exception:
        try:
            birth_date = datetime.strptime(str(birth_value), "%d/%m/%Y").date()
        except Exception:
            return None
    today = date.today()
    return int(today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day)))


def _format_birth_date(value) -> str:
    from datetime import datetime
    if not value:
        return "-"
    if hasattr(value, "strftime"):
        return value.strftime("%d/%m/%Y")
    value_str = str(value)
    try:
        return datetime.fromisoformat(value_str).strftime("%d/%m/%Y")
    except Exception:
        return value_str


def _build_pdf_header_html(patient: dict, profile: dict | None = None) -> str:
    birth_date = _format_birth_date(patient.get("birth_date"))
    age = _patient_age_from_birth_date(patient.get("birth_date"))
    age_text = f"{age} anos" if age is not None else "-"
    profile = profile or {}
    professional_name = (profile.get("full_name") or "").strip()
    professional_role = (profile.get("profession") or "").strip()
    professional_text = " / ".join(part for part in [professional_name, professional_role] if part) or "-"
    return (
        '<div class="pdf-document">'
        '<div style="font-size:17px; font-weight:700; color:#0f172a; text-align:center; text-transform:uppercase; letter-spacing:0.04em; margin-bottom:12px;">Relatório de Avaliação Neuropsicológica</div>'
        '<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse; font-size:10.5pt; margin-bottom:14px;">'
        f'<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; width:18%; font-size:10pt; font-weight:700; background:#e2e8f0;">Paciente</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; width:32%;">{escape(patient.get("full_name") or "Paciente")}</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; width:18%; font-size:10pt; font-weight:700; background:#e2e8f0;">Responsável</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; width:32%;">{escape(patient.get("responsavel") or "-")}</td>'
        f'</tr>'
        f'<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; font-size:10pt; font-weight:700; background:#e2e8f0;">Idade</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px;">{escape(age_text)}</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; font-size:10pt; font-weight:700; background:#e2e8f0;">Profissional</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px;">{escape(professional_text)}</td>'
        f'</tr>'
        f'<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; font-size:10pt; font-weight:700; background:#e2e8f0;">Data de nascimento</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px;">{escape(birth_date)}</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px; font-size:10pt; font-weight:700; background:#e2e8f0;">Gênero</td>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:7px 8px;">{escape(patient.get("gender") or "-")}</td>'
        f'</tr>'
        '</table>'
    )


def _build_patient_description_html(patient_description: str) -> str:
    if not patient_description or not patient_description.strip():
        return ""
    description_html = escape(patient_description).replace("\n", "<br />")
    return (
        '<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse; margin-bottom:14px;">'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; background:#f8fafc; font-size:11pt; font-weight:700;">Descrição do paciente</td>'
        '</tr>'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; font-size:10.5pt; line-height:1.55;">{description_html}</td>'
        '</tr>'
        '</table>'
    )


def _build_patient_history_html(patient_health_history: str) -> str:
    if not patient_health_history or not patient_health_history.strip():
        return ""
    history_html = escape(patient_health_history).replace("\n", "<br />")
    return (
        '<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse; margin-bottom:14px;">'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; background:#f8fafc; font-size:11pt; font-weight:700;">Histórico de saúde</td>'
        '</tr>'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; font-size:10.5pt; line-height:1.55;">{history_html}</td>'
        '</tr>'
        '</table>'
    )


def _build_patient_school_life_html(patient_school_life: str) -> str:
    if not patient_school_life or not patient_school_life.strip():
        return ""
    school_life_html = escape(patient_school_life).replace("\n", "<br />")
    return (
        '<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse; margin-bottom:14px;">'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; background:#f8fafc; font-size:11pt; font-weight:700;">Vida Escolar</td>'
        '</tr>'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; font-size:10.5pt; line-height:1.55;">{school_life_html}</td>'
        '</tr>'
        '</table>'
    )


def _build_patient_evaluation_behavior_html(patient_evaluation_behavior: str) -> str:
    if not patient_evaluation_behavior or not patient_evaluation_behavior.strip():
        return ""
    behavior_html = escape(patient_evaluation_behavior).replace("\n", "<br />")
    return (
        '<table width="100%" cellpadding="0" cellspacing="0" style="border-collapse:collapse; margin-bottom:14px;">'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; background:#f8fafc; font-size:11pt; font-weight:700;">Comportamento durante a avaliação</td>'
        '</tr>'
        '<tr>'
        f'<td style="border:0.5px solid {PDF_BORDER_COLOR}; padding:8px 10px; font-size:10.5pt; line-height:1.55;">{behavior_html}</td>'
        '</tr>'
        '</table>'
    )


def _build_pdf_page_html(body_html: str) -> str:
    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    @page {{ size: A4; margin: 14mm; }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      padding: 0;
      color: #111827;
      font-family: Arial, Helvetica, sans-serif;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
      background: #fff;
    }}
    .pdf-document {{
      font-size: 10.5pt;
      line-height: 1.45;
    }}
    .pdf-title {{
      font-size: 15pt;
      font-weight: 700;
      text-align: center;
      text-transform: uppercase;
      letter-spacing: 0.04em;
      margin: 0 0 12px 0;
      color: #0f172a;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
    }}
    .block-table {{
      margin-bottom: 14px;
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .block-table td {{
      border: 0.5px solid #cbd5e1;
      padding: 7px 8px;
      vertical-align: top;
    }}
    .label {{
      width: 18%;
      background: #e2e8f0;
      font-weight: 700;
    }}
    .section-title {{
      background: #e2e8f0;
      font-weight: 700;
      text-transform: uppercase;
    }}
    .section-body {{
      padding: 8px 10px;
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .section-body > * {{
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .report-title {{
      font-size: 11pt;
      font-weight: 700;
      margin: 0 0 8px 0;
      text-transform: uppercase;
      letter-spacing: 0.02em;
    }}
    .report-box {{
      margin-bottom: 12px;
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .report-box table {{
      font-size: 10pt;
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .report-box th, .report-box td {{
      border: 0.5px solid #cbd5e1;
      padding: 6px 8px;
    }}
    .report-box tr,
    .report-box thead,
    .report-box tbody {{
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .report-box thead th {{
      background: #e2e8f0;
      font-weight: 700;
    }}
    .report-box tbody tr:nth-child(even) td {{
      background: #f8fafc;
    }}
    .conclusion p {{
      margin: 0 0 8px 0;
      text-align: justify;
      font-size: 10.75pt;
      line-height: 1.6;
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .conclusion {{
      break-inside: avoid;
      page-break-inside: avoid;
    }}
  </style>
</head>
<body>
  {body_html}
</body>
</html>"""
SPA_INDEX = FRONTEND_DIST / "index.html"

@asynccontextmanager
async def lifespan(_app: FastAPI):
    ensure_report_folders()
    yield


app = FastAPI(title="Report Psicologia API", version="1.0.0", lifespan=lifespan)
app.add_middleware(
    SessionMiddleware,
    secret_key=SETTINGS["secret_key"],
    same_site="lax",
    https_only=False,
    max_age=None,
)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

# Serve o frontend React buildado (bundles hasheados do Vite)
if (FRONTEND_DIST / "assets").is_dir():
    app.mount("/assets", StaticFiles(directory=str(FRONTEND_DIST / "assets")), name="assets")


def serve_spa():
    """Devolve o index.html do SPA React (login/rotas tratadas no cliente)."""
    return FileResponse(str(SPA_INDEX))

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SUPABASE_URL = SETTINGS["supabase_url"]
SUPABASE_KEY = SETTINGS["supabase_key"]


@app.get("/health")
def health_check():
    return {"status": "ok", "supabase": "configured" if SUPABASE_URL and SUPABASE_KEY else "missing"}


@app.get("/")
async def root(request: Request):
    return serve_spa()


@app.get("/login")
async def login_page(request: Request):
    return serve_spa()


@app.post("/login")
async def login(request: Request, email: str = Form(...), password: str = Form(...)):
    try:
        await login_user(email, password, request)
    except Exception as exc:
        detail = str(exc)
        if 'Invalid login credentials' in detail:
            detail = 'E-mail ou senha invÃ¡lidos.'
        return templates.TemplateResponse(request, "login.html", {"request": request, "error": detail})

    return RedirectResponse(url="/generate-report", status_code=303)


@app.post("/api/auth/login")
async def api_login(request: Request):
    payload = await request.json()
    email = payload.get("email")
    password = payload.get("password")
    if not email or not password:
        raise HTTPException(status_code=400, detail="Email e senha sÃ£o obrigatÃ³rios")
    try:
        await login_user(email, password, request)
    except HTTPException as exc:
        raise exc
    except Exception as exc:
        detail = str(exc)
        if 'Invalid login credentials' in detail:
            detail = 'E-mail ou senha invÃ¡lidos.'
        raise HTTPException(status_code=401, detail=detail)
    return {"ok": True, "user": request.session.get("user")}


@app.post("/api/auth/logout")
async def api_logout(request: Request):
    await logout_user(request)
    return {"ok": True}


@app.get("/api/auth/user")
def api_user(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")
    return {"user": user}


@app.get("/logout")
async def logout(request: Request):
    await logout_user(request)
    return RedirectResponse(url="/login", status_code=303)


@app.get("/account")
async def account(request: Request):
    return serve_spa()


@app.post("/account/update-profile")
async def update_profile(request: Request, full_name: str = Form(""), profession: str = Form(""), gender: str = Form("")):
    user = get_user_from_session(request)
    if not user:
        return RedirectResponse(url="/login", status_code=303)

    full_name = (full_name or "").strip()
    profession = (profession or "").strip()
    gender = (gender or "").strip()

    allowed_genders = {"Masculino", "Feminino", "Outro"}
    if gender and gender not in allowed_genders:
        raise HTTPException(status_code=400, detail="GÃªnero invÃ¡lido")

    client = get_authenticated_client(request)
    response = client.table("profiles").update({
        "full_name": full_name or user.get("full_name") or None,
        "profession": profession or None,
        "gender": gender or None,
    }).eq("id", user["id"]).execute()

    if getattr(response, "error", None):
        raise HTTPException(status_code=500, detail=str(response.error))

    request.session['user'] = {
        **user,
        'full_name': full_name or user.get('full_name'),
        'profession': profession or user.get('profession'),
        'gender': gender or user.get('gender'),
    }

    return RedirectResponse(url="/account?updated=1", status_code=303)


@app.patch("/api/account")
async def api_update_account(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")

    payload = await request.json()
    full_name = (payload.get("full_name") or "").strip()
    profession = (payload.get("profession") or "").strip()
    gender = (payload.get("gender") or "").strip()

    allowed_genders = {"Masculino", "Feminino", "Outro"}
    if gender and gender not in allowed_genders:
        raise HTTPException(status_code=400, detail="GÃªnero invÃ¡lido")

    client = get_authenticated_client(request)
    response = client.table("profiles").update({
        "full_name": full_name or user.get("full_name") or None,
        "profession": profession or None,
        "gender": gender or None,
    }).eq("id", user["id"]).execute()

    if getattr(response, "error", None):
        raise HTTPException(status_code=500, detail=str(response.error))

    updated_user = {
        **user,
        'full_name': full_name or user.get('full_name'),
        'profession': profession or user.get('profession'),
        'gender': gender or user.get('gender'),
    }
    request.session['user'] = updated_user
    return {"ok": True, "user": updated_user}


@app.get("/generate-report")
async def generate_report(request: Request):
    return serve_spa()


@app.get("/api/patients")
async def api_patients(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")

    client = get_authenticated_client(request)
    # select additional fields so frontend can display gender and contact and pre-fill edit form
    response = client.table("patients").select("id, full_name, birth_date, gender, phone, email, responsavel").eq("psychologist_id", user["id"]).order("created_at", desc=True).execute()
    if getattr(response, "error", None):
        raise HTTPException(status_code=500, detail=str(response.error))

    raw_patients = getattr(response, "data", []) or []
    from datetime import date, datetime
    today = date.today()
    patients = []
    for p in raw_patients:
        age = ""
        birth_value = p.get("birth_date")
        birth_date_display = ""
        if birth_value:
            try:
                birth_date = datetime.fromisoformat(str(birth_value)).date()
                age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
                birth_date_display = birth_date.strftime("%d/%m/%Y")
            except Exception:
                age = ""
                birth_date_display = str(birth_value)

        patients.append({
            "id": p.get("id"),
            "full_name": p.get("full_name") or "Paciente",
            "age": age,
            "birth_date": birth_date_display,
            "gender": p.get("gender") or None,
            "phone": p.get("phone") or None,
            "email": p.get("email") or None,
            "responsavel": p.get("responsavel") or None,
        })
    return patients


@app.get("/api/report-types")
async def api_report_types(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")
    return get_report_folders()


@app.get("/api/report-fields")
async def api_report_fields(request: Request, report_name: str = ""):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")
    if not report_name or report_name not in get_report_folders():
        raise HTTPException(status_code=400, detail="RelatÃ³rio invÃ¡lido")
    return get_report_input_fields(report_name)


@app.post("/api/reports")
async def create_report(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")

    payload = await request.json()
    patient_id = payload.get("patient_id")
    report_name = payload.get("report_name")
    input_data = payload.get("input_data") or {}
    if not patient_id or not report_name:
        raise HTTPException(status_code=400, detail="patient_id e report_name sÃ£o obrigatÃ³rios")
    report_folders = get_report_folders()
    if report_name not in report_folders:
        raise HTTPException(status_code=400, detail="RelatÃ³rio invÃ¡lido")

    client = get_authenticated_client(request)
    patient_resp = client.table("patients").select("id, full_name, birth_date, responsavel").eq("psychologist_id", user["id"]).eq("id", patient_id).limit(1).execute()
    if getattr(patient_resp, "error", None):
        raise HTTPException(status_code=500, detail=str(patient_resp.error))
    raw_data = getattr(patient_resp, "data", []) or []
    if not raw_data:
        raise HTTPException(status_code=404, detail="Paciente nÃ£o encontrado")
    patient = raw_data[0]

    report_text = ""
    if report_name == "TAC 2":
        report_text = build_tac2_text_report(
            client,
            patient["id"],
            patient.get("full_name") or "Paciente",
            input_data,
        )
    else:
        report_module = None
        try:
            from app.report_store import load_report_module
            report_module = load_report_module(report_name)
        except Exception:
            report_module = None

        if not report_module or not hasattr(report_module, 'build_report'):
            raise HTTPException(status_code=400, detail="RelatÃ³rio nÃ£o suportado")

        # Report modules are called without the DB client, so they cannot
        # resolve the patient's age themselves. Inject the patient's age so the
        # report can select the correct age column in its score tables.
        report_input = dict(input_data or {})
        patient_age = _patient_age_from_birth_date(patient.get("birth_date"))
        if patient_age is not None:
            report_input["_patient_age"] = patient_age

        report_output = report_module.build_report(patient["id"], patient.get("full_name") or "Paciente", report_input)
        if isinstance(report_output, str):
            report_text = report_output

    if report_text:
        from app.gemini_service import generate_interpretation

        table_match = re.search(r"<table.*?>.*?</table>", report_text, re.DOTALL | re.IGNORECASE)
        table_html = table_match.group(0) if table_match else ""
        observations = input_data.get("observacoes_sobre_o_teste", "")
        interpretation = await generate_interpretation(report_name, observations, table_html, patient.get("birth_date"))
        ai_html = (
            '<div style="margin-top:20px;">'
            f'<div style="border:1px solid #cbd5e1; background:#f8fafc; padding:10px; border-radius:6px; font-size:9pt; line-height:1.5;">{interpretation}</div>'
            '</div>'
        )
        if report_text.endswith("</div>\n"):
            report_text = report_text[:-7] + ai_html + "</div>\n"
        elif report_text.endswith("</div>"):
            report_text = report_text[:-6] + ai_html + "</div>"
        else:
            report_text += ai_html

        return HTMLResponse(report_text)

    return {"ok": True, "report_name": report_name, "patient_id": patient_id}


async def build_report_html(client, patient, report_name: str, input_data: dict) -> str:
    report_text = ""
    if report_name == "TAC 2":
        report_text = build_tac2_text_report(
            client,
            patient["id"],
            patient.get("full_name") or "Paciente",
            input_data,
        )
    else:
        report_module = None
        try:
            from app.report_store import load_report_module
            report_module = load_report_module(report_name)
        except Exception:
            report_module = None

        if not report_module or not hasattr(report_module, 'build_report'):
            raise HTTPException(status_code=400, detail="RelatÃ³rio nÃ£o suportado")

        # Report modules are called without the DB client, so they cannot
        # resolve the patient's age themselves. Inject the patient's age so the
        # report can select the correct age column in its score tables.
        report_input = dict(input_data or {})
        patient_age = _patient_age_from_birth_date(patient.get("birth_date"))
        if patient_age is not None:
            report_input["_patient_age"] = patient_age

        report_output = report_module.build_report(patient["id"], patient.get("full_name") or "Paciente", report_input)
        if isinstance(report_output, str):
            report_text = report_output

    if not report_text:
        raise HTTPException(status_code=400, detail="RelatÃ³rio vazio")

    from app.gemini_service import generate_interpretation
    table_match = re.search(r"<table.*?>.*?</table>", report_text, re.DOTALL | re.IGNORECASE)
    table_html = table_match.group(0) if table_match else ""
    observations = input_data.get("observacoes_sobre_o_teste", "")
    interpretation = await generate_interpretation(report_name, observations, table_html, patient.get("birth_date"))
    ai_html = (
        '<div style="margin-top:20px;">'
        f'<div style="border:1px solid #cbd5e1; background:#f8fafc; padding:10px; border-radius:6px; font-size:9pt; line-height:1.5;">{interpretation}</div>'
        '</div>'
    )
    if report_text.endswith("</div>\n"):
        report_text = report_text[:-7] + ai_html + "</div>\n"
    elif report_text.endswith("</div>"):
        report_text = report_text[:-6] + ai_html + "</div>"
    else:
        report_text += ai_html

    return report_text


async def build_report_html_no_ai(client, patient, report_name: str, input_data: dict) -> str:
    report_text = ""
    if report_name == "TAC 2":
        report_text = build_tac2_text_report(
            client,
            patient["id"],
            patient.get("full_name") or "Paciente",
            input_data,
        )
    else:
        report_module = None
        try:
            from app.report_store import load_report_module
            report_module = load_report_module(report_name)
        except Exception:
            report_module = None

        if not report_module or not hasattr(report_module, 'build_report'):
            raise HTTPException(status_code=400, detail="RelatÃ³rio nÃ£o suportado")

        report_input = dict(input_data or {})
        patient_age = _patient_age_from_birth_date(patient.get("birth_date"))
        if patient_age is not None:
            report_input["_patient_age"] = patient_age

        report_output = report_module.build_report(patient["id"], patient.get("full_name") or "Paciente", report_input)
        if isinstance(report_output, str):
            report_text = report_output

    if not report_text:
        raise HTTPException(status_code=400, detail="RelatÃ³rio vazio")

    return report_text


def _render_pdf_sync(html_page: str) -> bytes:
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        try:
            page = browser.new_page(viewport={"width": 1240, "height": 1754})
            page.set_content(html_page, wait_until="networkidle")
            return page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "14mm", "right": "14mm", "bottom": "14mm", "left": "14mm"},
            )
        finally:
            browser.close()


async def build_combined_pdf(report_sections: list[str]) -> bytes:
    html_page = _build_pdf_page_html("".join(report_sections))
    return await asyncio.to_thread(_render_pdf_sync, html_page)


async def _build_conclusion_html(
    patient: dict,
    report_results: list,
    patient_description: str,
    patient_health_history: str,
    patient_school_life: str,
    patient_evaluation_behavior: str,
    user_ia_direction_conclusion: str,
) -> str:
    from app.conclusao_service import generate_conclusion

    conclusion_html = await generate_conclusion(
        birth_date=patient.get('birth_date'),
        report_results=report_results,
        patient_description=patient_description,
        patient_health_history=patient_health_history,
        patient_school_life=patient_school_life,
        patient_evaluation_behavior=patient_evaluation_behavior,
        user_ia_direction_conclusion=user_ia_direction_conclusion,
    )
    return (
        '<table class="block-table report-block" cellpadding="0" cellspacing="0">'
        '<tr>'
        '<td class="section-title">Síntese dos resultados</td>'
        '</tr>'
        '<tr>'
        f'<td class="section-body conclusion">{conclusion_html}</td>'
        '</tr>'
        '</table>'
    )


def _wrap_report_html(report_name: str, report_html: str) -> str:
    return (
        '<table class="block-table report-box report-block" cellpadding="0" cellspacing="0">'
        '<tr>'
        f'<td class="section-title">{escape(report_name or "Relatório")}</td>'
        '</tr>'
        '<tr>'
        f'<td class="section-body">{report_html}</td>'
        '</tr>'
        '</table>'
    )


class _SimpleHtmlBlockParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.blocks = []
        self._current = []
        self._current_tag = None

    def _flush(self):
        text = ''.join(self._current).strip()
        if text:
            self.blocks.append((self._current_tag or 'p', text))
        self._current = []
        self._current_tag = None

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in {'p', 'div', 'h1', 'h2', 'h3', 'h4', 'li', 'td', 'th', 'tr'}:
            if self._current:
                self._flush()
            self._current_tag = tag
        elif tag == 'br':
            self._current.append('\n')
        elif tag == 'table':
            if self._current:
                self._flush()
            self.blocks.append(('table_start', ''))

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in {'p', 'div', 'h1', 'h2', 'h3', 'h4', 'li', 'td', 'th', 'tr'}:
            self._flush()

    def handle_data(self, data):
        self._current.append(data)


def _html_to_docx_bytes(html_text: str) -> bytes:
    parser = _SimpleHtmlBlockParser()
    parser.feed(html_text or '')

    def mark_paragraph(paragraph, keep_together=True, keep_with_next=False):
        fmt = paragraph.paragraph_format
        fmt.keep_together = keep_together
        fmt.keep_with_next = keep_with_next
        fmt.widow_control = True

    def mark_row_no_split(row):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        tr_pr.append(cant_split)

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    previous_heading = None
    for tag, text in parser.blocks:
        cleaned = re.sub(r'\s+', ' ', text).strip()
        if not cleaned:
            continue
        if tag == 'h1':
            paragraph = doc.add_heading(cleaned, level=1)
            mark_paragraph(paragraph, keep_together=True, keep_with_next=True)
            previous_heading = paragraph
        elif tag == 'h2':
            paragraph = doc.add_heading(cleaned, level=2)
            mark_paragraph(paragraph, keep_together=True, keep_with_next=True)
            previous_heading = paragraph
        elif tag == 'h3':
            paragraph = doc.add_heading(cleaned, level=3)
            mark_paragraph(paragraph, keep_together=True, keep_with_next=True)
            previous_heading = paragraph
        elif tag == 'li':
            paragraph = doc.add_paragraph(cleaned, style='List Bullet')
            mark_paragraph(paragraph, keep_together=True)
        elif tag == 'table_start':
            paragraph = doc.add_paragraph('Tabela do relatório')
            mark_paragraph(paragraph, keep_together=True)
        else:
            paragraph = doc.add_paragraph(cleaned)
            mark_paragraph(paragraph, keep_together=True)

    for table in doc.tables:
        for row in table.rows:
            mark_row_no_split(row)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_editor_pdf_html(html_text: str) -> str:
    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    @page {{ size: A4; margin: 18mm; }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      padding: 0;
      color: #111827;
      font-family: Arial, Helvetica, sans-serif;
      background: #fff;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}
    .editor-document {{
      font-size: 11pt;
      line-height: 1.55;
    }}
    .editor-document h1 {{
      font-size: 16pt;
      margin: 0 0 12px 0;
    }}
    .editor-document h2 {{
      font-size: 14pt;
      margin: 0 0 10px 0;
    }}
    .editor-document h3 {{
      font-size: 12pt;
      margin: 0 0 8px 0;
    }}
    .editor-document p {{
      margin: 0 0 8px 0;
      text-align: justify;
    }}
    .editor-document table {{
      width: 100%;
      border-collapse: collapse;
      margin-bottom: 12px;
    }}
    .editor-document td, .editor-document th {{
      border: 0.5px solid #cbd5e1;
      padding: 6px 8px;
      vertical-align: top;
    }}
  </style>
</head>
<body>
  <div class="editor-document">{html_text}</div>
</body>
</html>"""


def _render_editor_pdf_sync(html_page: str) -> bytes:
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        try:
            page = browser.new_page(viewport={"width": 1240, "height": 1754})
            page.set_content(html_page, wait_until="networkidle")
            return page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "18mm", "right": "18mm", "bottom": "18mm", "left": "18mm"},
            )
        finally:
            browser.close()

@app.post('/api/reports/pdf')
async def create_reports_pdf(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail='NÃ£o autenticado')

    payload = await request.json()
    patient_id = payload.get('patient_id')
    report_entries = payload.get('report_entries') or []
    patient_description = payload.get('patient_description') or ''
    patient_health_history = payload.get('patient_health_history') or ''
    patient_school_life = payload.get('patient_school_life') or ''
    patient_evaluation_behavior = payload.get('patient_evaluation_behavior') or ''
    user_ia_direction_conclusion = payload.get('user_ia_direction_conclusion') or ''
    document_html = payload.get('document_html') or ''
    if not patient_id or not isinstance(report_entries, list) or not report_entries:
        raise HTTPException(status_code=400, detail='patient_id e report_entries sÃ£o obrigatÃ³rios')

    if document_html.strip():
        pdf_content = await build_combined_pdf([document_html])
        return StreamingResponse(io.BytesIO(pdf_content), media_type='application/pdf', headers={'Content-Disposition': 'inline; filename="Relat\u00f3rio.pdf"'})

    report_folders = get_report_folders()

    client = get_authenticated_client(request)
    patient_resp = client.table('patients').select('id, full_name, birth_date, gender, responsavel').eq('psychologist_id', user['id']).eq('id', patient_id).limit(1).execute()
    if getattr(patient_resp, 'error', None):
        raise HTTPException(status_code=500, detail=str(patient_resp.error))
    raw_data = getattr(patient_resp, 'data', []) or []
    if not raw_data:
        raise HTTPException(status_code=404, detail='Paciente nÃ£o encontrado')
    patient = raw_data[0]

    profile_resp = client.table('profiles').select('id, profession').eq('id', user['id']).limit(1).execute()
    if getattr(profile_resp, 'error', None):
        raise HTTPException(status_code=500, detail=str(profile_resp.error))
    profile_rows = getattr(profile_resp, 'data', []) or []
    profile = profile_rows[0] if profile_rows else {}

    report_sections = []
    report_results = []
    for entry in report_entries:
        report_name = entry.get('report_name')
        input_data = entry.get('input_data') or {}
        if not report_name or report_name not in report_folders:
            raise HTTPException(status_code=400, detail=f"RelatÃ³rio invÃ¡lido: {report_name}")
        report_html = await build_report_html(client, patient, report_name, input_data)
        report_sections.append(_wrap_report_html(report_name, report_html))
        report_results.append({'report_name': report_name, 'results_html': report_html})

    header_html = _build_pdf_header_html(patient, profile)
    description_html = _build_patient_description_html(patient_description)
    history_html = _build_patient_history_html(patient_health_history)
    school_life_html = _build_patient_school_life_html(patient_school_life)
    behavior_html = _build_patient_evaluation_behavior_html(patient_evaluation_behavior)
    conclusion_html = await _build_conclusion_html(
        patient,
        report_results,
        patient_description,
        patient_health_history,
        patient_school_life,
        patient_evaluation_behavior,
        user_ia_direction_conclusion,
    )
    body_html = ''.join(report_sections)
    pdf_content = await build_combined_pdf([header_html + description_html + history_html + school_life_html + behavior_html + body_html + conclusion_html])
    return StreamingResponse(io.BytesIO(pdf_content), media_type='application/pdf', headers={'Content-Disposition': 'inline; filename="Relat\u00f3rio.pdf"'})


@app.post('/api/reports/pdf-html')
async def create_reports_pdf_html(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail='N\u00e3o autenticado')

    payload = await request.json()
    patient_id = payload.get('patient_id')
    report_entries = payload.get('report_entries') or []
    patient_description = payload.get('patient_description') or ''
    patient_health_history = payload.get('patient_health_history') or ''
    patient_school_life = payload.get('patient_school_life') or ''
    patient_evaluation_behavior = payload.get('patient_evaluation_behavior') or ''
    user_ia_direction_conclusion = payload.get('user_ia_direction_conclusion') or ''
    if not patient_id or not isinstance(report_entries, list) or not report_entries:
        raise HTTPException(status_code=400, detail='patient_id e report_entries s\u00e3o obrigat\u00f3rios')

    report_folders = get_report_folders()

    client = get_authenticated_client(request)
    patient_resp = client.table('patients').select('id, full_name, birth_date, gender, responsavel').eq('psychologist_id', user['id']).eq('id', patient_id).limit(1).execute()
    if getattr(patient_resp, 'error', None):
        raise HTTPException(status_code=500, detail=str(patient_resp.error))
    raw_data = getattr(patient_resp, 'data', []) or []
    if not raw_data:
        raise HTTPException(status_code=404, detail='Paciente n\u00e3o encontrado')
    patient = raw_data[0]

    profile_resp = client.table('profiles').select('id, profession').eq('id', user['id']).limit(1).execute()
    if getattr(profile_resp, 'error', None):
        raise HTTPException(status_code=500, detail=str(profile_resp.error))
    profile_rows = getattr(profile_resp, 'data', []) or []
    profile = profile_rows[0] if profile_rows else {}

    report_sections = []
    report_results = []
    for entry in report_entries:
        report_name = entry.get('report_name')
        input_data = entry.get('input_data') or {}
        if not report_name or report_name not in report_folders:
            raise HTTPException(status_code=400, detail=f"Relatório inválido: {report_name}")
        report_html = await build_report_html(client, patient, report_name, input_data)
        report_sections.append(_wrap_report_html(report_name, report_html))
        report_results.append({'report_name': report_name, 'results_html': report_html})

    header_html = _build_pdf_header_html(patient, profile)
    description_html = _build_patient_description_html(patient_description)
    history_html = _build_patient_history_html(patient_health_history)
    school_life_html = _build_patient_school_life_html(patient_school_life)
    behavior_html = _build_patient_evaluation_behavior_html(patient_evaluation_behavior)
    conclusion_html = await _build_conclusion_html(
        patient,
        report_results,
        patient_description,
        patient_health_history,
        patient_school_life,
        patient_evaluation_behavior,
        user_ia_direction_conclusion,
    )
    body_html = ''.join(report_sections)
    document_html = header_html + description_html + history_html + school_life_html + behavior_html + body_html + conclusion_html
    return {"html": document_html}


@app.post('/api/reports/editor-html')
async def create_reports_editor_html(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail='N\u00e3o autenticado')

    payload = await request.json()
    patient_id = payload.get('patient_id')
    report_entries = payload.get('report_entries') or []
    patient_description = payload.get('patient_description') or ''
    patient_health_history = payload.get('patient_health_history') or ''
    patient_school_life = payload.get('patient_school_life') or ''
    patient_evaluation_behavior = payload.get('patient_evaluation_behavior') or ''
    document_html = payload.get('document_html') or ''
    if not patient_id or not isinstance(report_entries, list) or not report_entries:
        raise HTTPException(status_code=400, detail='patient_id e report_entries s\u00e3o obrigat\u00f3rios')

    if document_html.strip():
        docx_bytes = _html_to_docx_bytes(document_html)
        return {
            "html": document_html,
            "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
            "file_name": "relatorio-editavel.docx",
        }

    report_folders = get_report_folders()

    client = get_authenticated_client(request)
    patient_resp = client.table('patients').select('id, full_name, birth_date, gender, responsavel').eq('psychologist_id', user['id']).eq('id', patient_id).limit(1).execute()
    if getattr(patient_resp, 'error', None):
        raise HTTPException(status_code=500, detail=str(patient_resp.error))
    raw_data = getattr(patient_resp, 'data', []) or []
    if not raw_data:
        raise HTTPException(status_code=404, detail='Paciente n\u00e3o encontrado')
    patient = raw_data[0]

    profile_resp = client.table('profiles').select('id, profession').eq('id', user['id']).limit(1).execute()
    if getattr(profile_resp, 'error', None):
        raise HTTPException(status_code=500, detail=str(profile_resp.error))
    profile_rows = getattr(profile_resp, 'data', []) or []
    profile = profile_rows[0] if profile_rows else {}

    report_sections = []
    report_results = []
    for entry in report_entries:
        report_name = entry.get('report_name')
        input_data = entry.get('input_data') or {}
        if not report_name or report_name not in report_folders:
            raise HTTPException(status_code=400, detail=f'Relat\u00f3rio inv\u00e1lido: {report_name}')
        report_html = await build_report_html_no_ai(client, patient, report_name, input_data)
        report_sections.append(_wrap_report_html(report_name, report_html))
        report_results.append({'report_name': report_name, 'results_html': report_html})

    header_html = _build_pdf_header_html(patient, profile)
    description_html = _build_patient_description_html(patient_description)
    history_html = _build_patient_history_html(patient_health_history)
    school_life_html = _build_patient_school_life_html(patient_school_life)
    behavior_html = _build_patient_evaluation_behavior_html(patient_evaluation_behavior)
    conclusion_html = ""
    body_html = ''.join(report_sections)
    document_html = header_html + description_html + history_html + school_life_html + behavior_html + body_html + conclusion_html
    docx_bytes = _html_to_docx_bytes(document_html)
    return {
        "html": document_html,
        "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
        "file_name": "relatorio-editavel.docx",
    }


@app.post('/api/reports/editor-pdf')
async def create_reports_editor_pdf(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail='N\u00e3o autenticado')

    payload = await request.json()
    html_text = payload.get('html') or ''
    if not html_text.strip():
        raise HTTPException(status_code=400, detail='html \u00e9 obrigat\u00f3rio')

    pdf_html = _build_editor_pdf_html(html_text)
    pdf_bytes = await asyncio.to_thread(_render_editor_pdf_sync, pdf_html)
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type='application/pdf',
        headers={'Content-Disposition': 'inline; filename="relatorio-editavel.pdf"'},
    )


@app.post('/api/conclusion')
async def api_conclusion(request: Request):
    """
    Gera a conclusÃ£o final do relatÃ³rio PDF usando conclusao_service.py.

    Recebe do frontend a idade (via data de nascimento do paciente), os
    resultados de cada relatÃ³rio selecionado, a descriÃ§Ã£o do paciente e a
    variÃ¡vel "user_ia_direction_conclusion".
    """
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail='NÃ£o autenticado')

    payload = await request.json()
    patient_id = payload.get('patient_id')
    report_results = payload.get('report_results') or []
    patient_description = payload.get('patient_description') or ''
    user_ia_direction_conclusion = payload.get('user_ia_direction_conclusion') or ''
    if not patient_id:
        raise HTTPException(status_code=400, detail='patient_id Ã© obrigatÃ³rio')

    client = get_authenticated_client(request)
    patient_resp = client.table('patients').select('id, full_name, birth_date').eq('psychologist_id', user['id']).eq('id', patient_id).limit(1).execute()
    if getattr(patient_resp, 'error', None):
        raise HTTPException(status_code=500, detail=str(patient_resp.error))
    raw_data = getattr(patient_resp, 'data', []) or []
    if not raw_data:
        raise HTTPException(status_code=404, detail='Paciente nÃ£o encontrado')
    patient = raw_data[0]

    from app.conclusao_service import generate_conclusion
    patient_health_history = payload.get('patient_health_history') or ''
    patient_school_life = payload.get('patient_school_life') or ''
    patient_evaluation_behavior = payload.get('patient_evaluation_behavior') or ''
    conclusion_html = await generate_conclusion(
        birth_date=patient.get('birth_date'),
        report_results=report_results,
        patient_description=patient_description,
        patient_health_history=patient_health_history,
        patient_school_life=patient_school_life,
        patient_evaluation_behavior=patient_evaluation_behavior,
        user_ia_direction_conclusion=user_ia_direction_conclusion,
    )
    return HTMLResponse(conclusion_html)


@app.get("/register-patient")
async def register_patient(request: Request):
    return serve_spa()


@app.get("/patients")
async def patients(request: Request):
    return serve_spa()


@app.patch("/api/patients/{patient_id}")
async def update_patient(patient_id: str, request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")

    if not request.session.get("access_token"):
        raise HTTPException(status_code=401, detail="UsuÃ¡rio nÃ£o autenticado")

    payload = await request.json()
    full_name = (payload.get("full_name") or "").strip()
    if not full_name:
        raise HTTPException(status_code=400, detail="Nome completo Ã© obrigatÃ³rio")

    def parse_birth_date(val):
        if not val:
            return None
        val = str(val).strip()
        from datetime import datetime
        try:
            if '/' in val:
                dt = datetime.strptime(val, '%d/%m/%Y')
            else:
                dt = datetime.strptime(val, '%Y-%m-%d')
            return dt.date().isoformat()
        except Exception:
            return None

    def normalize_phone(val):
        if not val:
            return None
        digits = ''.join(ch for ch in str(val) if ch.isdigit())
        if len(digits) < 10:
            return None
        aa = digits[:2]
        rest = digits[2:]
        if len(rest) == 8:
            return f"({aa}) {rest[:4]}-{rest[4:]}"
        return f"({aa}) {rest[:5]}-{rest[5:]}"

    birth_date = parse_birth_date(payload.get("birth_date") or None)
    if payload.get("birth_date") and birth_date is None:
        raise HTTPException(status_code=400, detail="Data de nascimento invÃ¡lida. Use o formato dd/mm/aaaa.")
    if birth_date:
        from datetime import date, datetime
        bd = datetime.fromisoformat(birth_date).date()
        age = date.today().year - bd.year - ((date.today().month, date.today().day) < (bd.month, bd.day))
        if age < 0:
            raise HTTPException(status_code=400, detail="Data de nascimento nÃ£o pode estar no futuro.")
        if age > 120:
            raise HTTPException(status_code=400, detail="A idade nÃ£o pode ser maior que 120 anos.")
    phone = normalize_phone(payload.get("phone") or None)
    gender = payload.get("gender") or None
    responsavel = (payload.get("responsavel") or "").strip() or None
    allowed_genders = {"Masculino", "Feminino", "Outro"}
    if gender and gender not in allowed_genders:
        raise HTTPException(status_code=400, detail="GÃªnero invÃ¡lido")

    data = {
        "full_name": full_name,
        "birth_date": birth_date,
        "gender": gender,
        "responsavel": responsavel,
        "phone": phone,
        "email": payload.get("email") or None,
    }

    client = get_authenticated_client(request)
    response = client.table("patients").update(data).eq("id", patient_id).eq("psychologist_id", user["id"]).execute()
    if getattr(response, "error", None):
        raise HTTPException(status_code=500, detail=str(response.error))

    updated_data = getattr(response, "data", []) or []
    if not updated_data:
        raise HTTPException(status_code=404, detail="Paciente nÃ£o encontrado")

    return {"ok": True, "patient": updated_data[0]}


@app.delete("/api/patients/{patient_id}")
async def delete_patient(patient_id: str, request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")

    if not request.session.get("access_token"):
        raise HTTPException(status_code=401, detail="UsuÃ¡rio nÃ£o autenticado")

    client = get_authenticated_client(request)
    response = client.table("patients").delete().eq("id", patient_id).eq("psychologist_id", user["id"]).execute()
    if getattr(response, "error", None):
        raise HTTPException(status_code=500, detail=str(response.error))

    deleted = getattr(response, "data", None) or []
    if not deleted:
        raise HTTPException(status_code=404, detail="Paciente nÃ£o encontrado")

    return {"ok": True, "deleted_id": patient_id}


@app.post("/api/patients")
async def create_patient(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")

    if not request.session.get("access_token"):
        raise HTTPException(status_code=401, detail="UsuÃ¡rio nÃ£o autenticado")

    payload = await request.json()
    full_name = (payload.get("full_name") or "").strip()
    if not full_name:
        raise HTTPException(status_code=400, detail="Nome completo Ã© obrigatÃ³rio")
    # check duplicate patient name for this psychologist
    client = get_authenticated_client(request)
    try:
        dup_check = client.table("patients").select("id").eq("psychologist_id", user["id"]).eq("full_name", full_name).limit(1).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    existing = getattr(dup_check, "data", None) or []
    if isinstance(existing, list) and existing:
        raise HTTPException(status_code=400, detail="JÃ¡ existe um paciente com esse nome registrado")

    # normalize and convert types according to supabase schema
    def parse_birth_date(val):
        if not val:
            return None
        val = str(val).strip()
        # accept dd/mm/yyyy or yyyy-mm-dd
        from datetime import datetime
        try:
            if '/' in val:
                dt = datetime.strptime(val, '%d/%m/%Y')
            else:
                dt = datetime.strptime(val, '%Y-%m-%d')
            return dt.date().isoformat()
        except Exception:
            return None

    def normalize_phone(val):
        if not val:
            return None
        digits = ''.join(ch for ch in str(val) if ch.isdigit())
        if len(digits) < 10:
            return None
        # area code first two digits
        aa = digits[:2]
        rest = digits[2:]
        if len(rest) == 8:
            # format (AA) 9999-9999
            return f"({aa}) {rest[:4]}-{rest[4:]}"
        else:
            # assume 9+4 => (AA) 99999-9999
            return f"({aa}) {rest[:5]}-{rest[5:]}"

    birth_date = parse_birth_date(payload.get("birth_date") or None)
    if payload.get("birth_date") and birth_date is None:
        raise HTTPException(status_code=400, detail="Data de nascimento invÃ¡lida. Use o formato dd/mm/aaaa.")
    if birth_date:
        from datetime import date, datetime
        bd = datetime.fromisoformat(birth_date).date()
        age = date.today().year - bd.year - ((date.today().month, date.today().day) < (bd.month, bd.day))
        if age < 0:
            raise HTTPException(status_code=400, detail="Data de nascimento nÃ£o pode estar no futuro.")
        if age > 120:
            raise HTTPException(status_code=400, detail="A idade nÃ£o pode ser maior que 120 anos.")
    phone = normalize_phone(payload.get("phone") or None)
    gender = payload.get("gender") or None
    responsavel = (payload.get("responsavel") or "").strip() or None
    allowed_genders = {"Masculino", "Feminino", "Outro"}
    if gender and gender not in allowed_genders:
        raise HTTPException(status_code=400, detail="GÃªnero invÃ¡lido")

    data = {
        "psychologist_id": user["id"],
        "full_name": full_name,
        "birth_date": birth_date,
        "gender": gender,
        "responsavel": responsavel,
        "phone": phone,
        "email": payload.get("email") or None,
    }

    response = client.table("patients").insert(data).execute()
    if getattr(response, "error", None):
        raise HTTPException(status_code=500, detail=str(response.error))

    return {"ok": True, "patient": getattr(response, "data", [None])[0]}


@app.get("/profiles")
def get_profiles(request: Request):
    client = get_authenticated_client(request)
    response = client.table("profiles").select("id, email, full_name").execute()
    if getattr(response, "error", None):
        raise HTTPException(status_code=500, detail=str(response.error))
    return getattr(response, "data", [])


@app.get("/chat-gemini")
async def chat_gemini_page(request: Request):
    return serve_spa()


@app.post("/api/chat-gemini")
async def api_chat_gemini(request: Request):
    user = get_user_from_session(request)
    if not user:
        raise HTTPException(status_code=401, detail="NÃ£o autenticado")

    payload = await request.json()
    contents = payload.get("contents")
    if not contents:
        raise HTTPException(status_code=400, detail="HistÃ³rico de mensagens Ã© obrigatÃ³rio")

    api_key = os.getenv("GEMINI_API_KEY")
    model = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
    if not api_key:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY nÃ£o configurada no servidor")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(url, json={"contents": contents}, timeout=60.0)
            response.raise_for_status()
            data = response.json()
            text = data['candidates'][0]['content']['parts'][0]['text']
            return {"text": text}
        except httpx.HTTPStatusError as e:
            try:
                err_detail = e.response.json()
            except Exception:
                err_detail = e.response.text
            raise HTTPException(status_code=e.response.status_code, detail=f"Erro na API do Gemini: {err_detail}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))


@app.get("/{path:path}")
async def spa_fallback(path: str):
    """Fallback do SPA: serve o index.html para rotas do cliente (React Router).

    /api, /static e /assets tÃªm rotas prÃ³prias; se chegarem aqui, devolvem 404.
    """
    if path.startswith(("api/", "static/", "assets/")):
        raise HTTPException(status_code=404, detail="Not found")
    return serve_spa()



